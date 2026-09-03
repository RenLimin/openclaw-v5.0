#!/usr/bin/env python3
"""
销售合同生成器
组件: SCA-001 (L4)
功能: 基于模板 + 变量填充生成合同 docx
依赖: python-docx
"""

import argparse
import os
import sqlite3
import sys
from datetime import datetime

DB_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "..", "contracts.db")

# 金额大写映射
DIGITS = "零壹贰叁肆伍陆柒捌玖"
UNITS = ["", "拾", "佰", "仟", "万", "拾", "佰", "仟", "亿"]


def amount_to_chinese(amount):
    """金额转中文大写"""
    if amount == 0:
        return "零元整"

    int_part = int(amount)
    dec_part = round((amount - int_part) * 100)

    # 整数部分
    int_str = str(int_part)
    result = ""
    zero_flag = False
    for i, ch in enumerate(int_str):
        digit = int(ch)
        pos = len(int_str) - 1 - i
        if digit == 0:
            zero_flag = True
        else:
            if zero_flag:
                result += "零"
                zero_flag = False
            result += DIGITS[digit] + UNITS[pos]

    # 处理末尾的万/亿单位
    if result.endswith("零"):
        result = result[:-1]
    result += "元"

    # 小数部分
    jiao = dec_part // 10
    fen = dec_part % 10
    if jiao == 0 and fen == 0:
        result += "整"
    else:
        if jiao > 0:
            result += DIGITS[jiao] + "角"
        if fen > 0:
            result += DIGITS[fen] + "分"

    return result


def get_contract_data(contract_id):
    """从数据库获取合同数据"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    contract = conn.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,)).fetchone()
    conn.close()
    return dict(contract) if contract else None


def generate_with_docx(args):
    """使用 python-docx 生成合同"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("❌ 需要安装 python-docx: pip install python-docx")
        sys.exit(1)

    data = get_contract_data(args.contract_id)
    if not data:
        print(f"❌ 合同 ID {args.contract_id} 不存在")
        return

    amount_cn = amount_to_chinese(data["amount"])

    # 创建文档
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data["title"])
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"

    # 合同编号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"合同编号：{data['contract_no']}").font.size = Pt(10)

    doc.add_paragraph()  # 空行

    # 双方信息
    p = doc.add_paragraph()
    p.add_run(f"甲方（委托方）：{data['party_a']}").bold = True
    p = doc.add_paragraph()
    p.add_run(f"地址：{data.get('party_a_address', '___________')}")
    p = doc.add_paragraph()
    p.add_run(f"联系电话：{data.get('party_a_phone', '___________')}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run(f"乙方（受托方）：{data['party_b']}").bold = True
    p = doc.add_paragraph()
    p.add_run(f"地址：{data.get('party_b_address', '___________')}")
    p = doc.add_paragraph()
    p.add_run(f"联系电话：{data.get('party_b_phone', '___________')}")

    doc.add_paragraph()

    # 前言
    p = doc.add_paragraph()
    p.add_run(
        f"本合同甲方委托乙方就 {data['title']} 项目进行的专项技术服务，"
        f"并支付相应的技术服务报酬。双方经过平等协商，在真实、充分地表达各自意愿的基础上，"
        f"根据《中华人民共和国民法典》的规定，达成如下协议，并由双方共同恪守。"
    )

    # 第一条
    doc.add_heading("第一条：技术服务的内容", level=2)
    p = doc.add_paragraph()
    p.add_run(f"1. 技术服务的目标：{data['title']}")
    p = doc.add_paragraph()
    p.add_run("2. 技术服务的内容：详见双方约定。")
    p = doc.add_paragraph()
    p.add_run("3. 技术服务的方式：远程交付。")

    # 第二条
    doc.add_heading("第二条：服务期限", level=2)
    p = doc.add_paragraph()
    p.add_run(
        f"服务期限：{data.get('effective_date', '____年__月__日')} 至 "
        f"{data.get('expiry_date', '____年__月__日')}。"
    )

    # 第三条
    doc.add_heading("第三条：技术服务报酬", level=2)
    p = doc.add_paragraph()
    p.add_run(
        f"1. 技术服务费总额为：¥{data['amount']:,.2f}（大写人民币：{amount_cn}）。"
    )
    p = doc.add_paragraph()
    p.add_run(
        f"2. 支付方式：由甲方于技术成果交付验收合格后 30 个工作日内一次性支付乙方。"
    )
    p = doc.add_paragraph()
    p.add_run(
        f"3. 乙方须出具与本合同金额相符的等额增值税专用发票（发票税率 {int(data.get('tax_rate', 0.06)*100)}%）。"
    )

    # 第四条 保密
    doc.add_heading("第四条：保密条款", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "甲乙双方应严格遵守《中华人民共和国保守国家秘密法》及国家相关保密工作的法律法规。"
        "未经信息披露方书面事先同意，不得向任何第三方披露保密资料。"
    )

    # 第五条 违约责任
    doc.add_heading("第五条：违约责任", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "1. 乙方延迟交付超过 10 个工作日的，每逾期一日应当承担迟延交付部分合同价款 1% 的违约金。"
    )
    p = doc.add_paragraph()
    p.add_run(
        "2. 甲方延迟支付合同价款超过 10 个工作日的，每逾期一日应当按延迟支付价款的 1% 承担违约金，"
        "最高不超过合同总额的 5%。"
    )

    # 第六条 争议解决
    doc.add_heading("第六条：争议解决", level=2)
    p = doc.add_paragraph()
    p.add_run("双方因履行本合同而发生的争议，应协商、调解解决。协商、调解不成的，依法向甲方所在地人民法院起诉。")

    # 第七条 其他
    doc.add_heading("第七条：其他", level=2)
    p = doc.add_paragraph()
    p.add_run(
        f"本合同一式肆份，甲方持贰份，乙方持贰份，具有同等法律效力。"
        f"本合同经双方签字盖章后生效。"
    )

    # 签署区
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run(f"甲方（盖章）：{data['party_a']}")
    p = doc.add_paragraph()
    p.add_run("法定代表人（或委托代理人）签字：")
    p = doc.add_paragraph()
    p.add_run("签字日期：    年    月    日")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run(f"乙方（盖章）：{data['party_b']}")
    p = doc.add_paragraph()
    p.add_run("法定代表人（或委托代理人）签字：")
    p = doc.add_paragraph()
    p.add_run("签字日期：    年    月    日")

    # 保存
    output_dir = os.path.join(os.path.dirname(__file__), "..", "..", "..", "contracts_output")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{data['contract_no']}.docx")
    doc.save(output_path)

    print(f"✅ 合同生成成功")
    print(f"   文件路径: {output_path}")
    print(f"   合同编号: {data['contract_no']}")
    print(f"   合同金额: ¥{data['amount']:,.2f}（{amount_cn}）")

    return output_path


def main():
    parser = argparse.ArgumentParser(description="销售合同生成器")
    sub = parser.add_subparsers(dest="command")

    # generate
    p_gen = sub.add_parser("generate", help="生成合同 docx")
    p_gen.add_argument("--contract-id", type=int, required=True, help="合同 ID")

    # amount to chinese
    p_amt = sub.add_parser("amount-cn", help="金额转大写")
    p_amt.add_argument("--amount", type=float, required=True)

    args = parser.parse_args()

    if args.command == "generate":
        generate_with_docx(args)
    elif args.command == "amount-cn":
        result = amount_to_chinese(args.amount)
        print(f"¥{args.amount:,.2f} → {result}")
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
