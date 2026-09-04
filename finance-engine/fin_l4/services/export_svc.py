"""报表导出服务 — Excel + Word"""

import io
from datetime import date
from decimal import Decimal
from typing import Dict, List, Optional

from fin_l4.db.repositories import (
    AccountRepository, TransactionRepository, LoanRepository,
    InsuranceRepository, PortfolioRepository, HoldingRepository,
)


class ExportService:
    """报表导出服务"""

    def __init__(self, conn):
        self.conn = conn
        self.account_repo = AccountRepository(conn)
        self.txn_repo = TransactionRepository(conn)
        self.loan_repo = LoanRepository(conn)
        self.insurance_repo = InsuranceRepository(conn)
        self.portfolio_repo = PortfolioRepository(conn)
        self.holding_repo = HoldingRepository(conn)

    # ── Excel 导出 ──

    def export_balance_sheet_excel(self, family_id: str) -> bytes:
        """资产负债表 → Excel"""
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

        wb = Workbook()
        ws = wb.active
        ws.title = "资产负债表"

        accounts = self.account_repo.list_by_family(family_id)

        # 样式
        header_font = Font(bold=True, size=12)
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font_white = Font(bold=True, size=12, color="FFFFFF")
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        # 标题
        ws.merge_cells('A1:C1')
        ws['A1'] = f"资产负债表 — {date.today()}"
        ws['A1'].font = Font(bold=True, size=14)

        # 资产
        row = 3
        ws.cell(row=row, column=1, value="资产").font = header_font
        row += 1
        ws.cell(row=row, column=1, value="科目").font = header_font_white
        ws.cell(row=row, column=1).fill = header_fill
        ws.cell(row=row, column=2, value="代码").font = header_font_white
        ws.cell(row=row, column=2).fill = header_fill
        ws.cell(row=row, column=3, value="余额").font = header_font_white
        ws.cell(row=row, column=3).fill = header_fill
        row += 1

        total_assets = Decimal("0")
        for acc in accounts:
            if acc["type"] != "ASSET":
                continue
            balance = self.account_repo.get_balance(acc["id"])
            if balance == 0:
                continue
            ws.cell(row=row, column=1, value=acc["name"]).border = thin_border
            ws.cell(row=row, column=2, value=acc["code"]).border = thin_border
            ws.cell(row=row, column=3, value=float(balance)).border = thin_border
            total_assets += balance
            row += 1

        ws.cell(row=row, column=1, value="资产合计").font = header_font
        ws.cell(row=row, column=3, value=float(total_assets)).font = header_font
        row += 2

        # 负债
        ws.cell(row=row, column=1, value="负债").font = header_font
        row += 1
        ws.cell(row=row, column=1, value="科目").font = header_font_white
        ws.cell(row=row, column=1).fill = header_fill
        ws.cell(row=row, column=2, value="代码").font = header_font_white
        ws.cell(row=row, column=2).fill = header_fill
        ws.cell(row=row, column=3, value="余额").font = header_font_white
        ws.cell(row=row, column=3).fill = header_fill
        row += 1

        total_liabilities = Decimal("0")
        for acc in accounts:
            if acc["type"] != "LIABILITY":
                continue
            balance = self.account_repo.get_balance(acc["id"])
            if balance == 0:
                continue
            ws.cell(row=row, column=1, value=acc["name"]).border = thin_border
            ws.cell(row=row, column=2, value=acc["code"]).border = thin_border
            ws.cell(row=row, column=3, value=float(balance)).border = thin_border
            total_liabilities += balance
            row += 1

        ws.cell(row=row, column=1, value="负债合计").font = header_font
        ws.cell(row=row, column=3, value=float(total_liabilities)).font = header_font
        row += 2

        # 净资产
        ws.cell(row=row, column=1, value="净资产").font = header_font
        ws.cell(row=row, column=3, value=float(total_assets - total_liabilities)).font = header_font

        # 列宽
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 18

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def export_transactions_excel(self, family_id: str,
                                  from_date: str = None,
                                  to_date: str = None) -> bytes:
        """交易明细 → Excel"""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Border, Side

        wb = Workbook()
        ws = wb.active
        ws.title = "交易明细"

        # 标题
        ws.merge_cells('A1:F1')
        ws['A1'] = f"交易明细 — {from_date or '全部'} ~ {to_date or '至今'}"
        ws['A1'].font = Font(bold=True, size=14)

        # 表头
        headers = ["日期", "借方科目", "贷方科目", "金额", "备注", "来源"]
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border

        # 数据
        sql = "SELECT * FROM fin4_transactions WHERE family_id = ?"
        params: list = [family_id]
        if from_date:
            sql += " AND date >= ?"
            params.append(from_date)
        if to_date:
            sql += " AND date <= ?"
            params.append(to_date)
        sql += " ORDER BY date DESC"

        rows = self.conn.execute(sql, params).fetchall()

        for i, r in enumerate(rows, 4):
            ws.cell(row=i, column=1, value=r["date"]).border = thin_border
            ws.cell(row=i, column=2, value=r.get("debit_account_id", "")).border = thin_border
            ws.cell(row=i, column=3, value=r.get("credit_account_id", "")).border = thin_border
            ws.cell(row=i, column=4, value=float(r["amount"])).border = thin_border
            ws.cell(row=i, column=5, value=r.get("note", "")).border = thin_border
            ws.cell(row=i, column=6, value=r.get("source", "")).border = thin_border

        # 列宽
        for col_letter, width in zip("ABCDEF", [12, 16, 16, 14, 20, 10]):
            ws.column_dimensions[col_letter].width = width

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def export_loan_schedule_excel(self, loan_data: Dict,
                                    schedule: List[Dict]) -> bytes:
        """还款计划 → Excel"""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Border, Side

        wb = Workbook()
        ws = wb.active
        ws.title = "还款计划"

        # 贷款信息
        ws['A1'] = f"贷款: {loan_data.get('name', '')}"
        ws['A1'].font = Font(bold=True, size=14)
        ws['A2'] = f"本金: {loan_data.get('principal', '')} | 利率: {loan_data.get('annual_rate', '')} | 期数: {loan_data.get('term_months', '')}"

        # 表头
        headers = ["期数", "月供", "本金", "利息", "剩余本金"]
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border

        for i, entry in enumerate(schedule, 5):
            ws.cell(row=i, column=1, value=entry["period"]).border = thin_border
            ws.cell(row=i, column=2, value=float(entry["payment"])).border = thin_border
            ws.cell(row=i, column=3, value=float(entry["principal"])).border = thin_border
            ws.cell(row=i, column=4, value=float(entry["interest"])).border = thin_border
            ws.cell(row=i, column=5, value=float(entry["remaining_balance"])).border = thin_border

        for col_letter, width in zip("ABCDE", [8, 14, 14, 14, 16]):
            ws.column_dimensions[col_letter].width = width

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    # ── Word 导出 ──

    def export_financial_report_word(self, family_id: str,
                                      health_result: Dict = None) -> bytes:
        """财务报告 → Word"""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 标题
        title = doc.add_heading("家庭财务分析报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"生成日期: {date.today()}")
        doc.add_paragraph("")

        # 资产负债表
        accounts = self.account_repo.list_by_family(family_id)

        doc.add_heading("一、资产负债表", level=1)

        asset_table = doc.add_table(rows=1, cols=3)
        asset_table.style = 'Light Grid Accent 1'
        hdr = asset_table.rows[0].cells
        hdr[0].text = "科目"
        hdr[1].text = "代码"
        hdr[2].text = "余额"

        total_assets = Decimal("0")
        total_liabilities = Decimal("0")

        for acc in accounts:
            balance = self.account_repo.get_balance(acc["id"])
            if balance == 0:
                continue
            row = asset_table.add_row().cells
            row[0].text = acc["name"]
            row[1].text = acc["code"]
            row[2].text = f"{balance:,.2f}"
            if acc["type"] == "ASSET":
                total_assets += balance
            elif acc["type"] == "LIABILITY":
                total_liabilities += balance

        doc.add_paragraph("")
        doc.add_paragraph(f"资产合计: {total_assets:,.2f}")
        doc.add_paragraph(f"负债合计: {total_liabilities:,.2f}")
        doc.add_paragraph(f"净资产: {total_assets - total_liabilities:,.2f}")

        # 贷款概览
        loans = self.loan_repo.list_by_family(family_id)
        if loans:
            doc.add_heading("二、贷款概览", level=1)
            for loan in loans:
                doc.add_paragraph(
                    f"• {loan['name']}: 本金 {loan['principal']}, "
                    f"利率 {float(loan['annual_rate'])*100:.2f}%, "
                    f"期数 {loan['term_months']}",
                    style='List Bullet'
                )

        # 保险概览
        policies = self.insurance_repo.list_by_family(family_id)
        if policies:
            doc.add_heading("三、保险概览", level=1)
            for pol in policies:
                doc.add_paragraph(
                    f"• {pol['product_name']}: 保额 {pol['sum_assured']}, "
                    f"年缴 {pol['annual_premium']}",
                    style='List Bullet'
                )

        # 理财建议
        if health_result:
            doc.add_heading("四、财务健康诊断", level=1)
            doc.add_paragraph(f"健康评分: {health_result.get('health_score', 'N/A')}")
            doc.add_paragraph(health_result.get("summary", ""))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
